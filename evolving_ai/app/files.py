from __future__ import annotations

import base64
import csv
from dataclasses import dataclass
from io import BytesIO, StringIO
from pathlib import Path
import re

from .attachments import Attachment


@dataclass(frozen=True, slots=True)
class ParsedFile:
    attachment: Attachment
    summary: str


class FileParser:
    def parse(self, *, filename: str, mime_type: str, payload: bytes) -> ParsedFile:
        suffix = Path(filename).suffix.lower()
        if mime_type.startswith("image/"):
            data_url = self._image_to_data_url(mime_type, payload)
            attachment = Attachment(
                name=filename,
                mime_type=mime_type,
                content=data_url,
                kind="image",
            )
            return ParsedFile(attachment=attachment, summary=f"{filename} image ready")

        if suffix == ".pdf" or mime_type == "application/pdf":
            content = self._parse_pdf(payload)
        elif suffix == ".docx" or mime_type in {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        }:
            content = self._parse_docx(payload)
        elif suffix == ".csv" or mime_type in {"text/csv", "application/csv"}:
            content = self._parse_csv(payload)
        elif suffix in {".html", ".htm"} or mime_type == "text/html":
            content = self._parse_html(payload)
        else:
            content = payload.decode("utf-8", errors="ignore")

        compact = content.strip() or f"No text could be extracted from {filename}."
        attachment = Attachment(
            name=filename,
            mime_type=mime_type or "text/plain",
            content=compact[:24000],
            kind="text",
        )
        return ParsedFile(
            attachment=attachment,
            summary=f"{filename} parsed into {len(attachment.content)} characters",
        )

    def _parse_pdf(self, payload: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ModuleNotFoundError:
            return "PDF parsing dependency is unavailable in this runtime."
        reader = PdfReader(BytesIO(payload))
        pieces: list[str] = []
        for page_index, page in enumerate(reader.pages[:12]):
            text = page.extract_text() or ""
            if text.strip():
                pieces.append(f"Page {page_index + 1}\n{text.strip()}")
        return "\n\n".join(pieces)

    def _parse_docx(self, payload: bytes) -> str:
        try:
            from docx import Document
        except ModuleNotFoundError:
            return "DOCX parsing dependency is unavailable in this runtime."
        document = Document(BytesIO(payload))
        parts: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    parts.append(" | ".join(values))
        return "\n".join(parts)

    def _parse_csv(self, payload: bytes) -> str:
        decoded = payload.decode("utf-8", errors="ignore")
        reader = csv.reader(StringIO(decoded))
        rows = list(reader)
        if not rows:
            return ""
        preview_rows = rows[:12]
        rendered = [" | ".join(row) for row in preview_rows]
        return "\n".join(rendered)

    def _parse_html(self, payload: bytes) -> str:
        decoded = payload.decode("utf-8", errors="ignore")
        try:
            from bs4 import BeautifulSoup
        except ModuleNotFoundError:
            without_hidden = re.sub(
                r"<(?:script|style|noscript)\b[^>]*>.*?</(?:script|style|noscript)>",
                " ",
                decoded,
                flags=re.IGNORECASE | re.DOTALL,
            )
            without_tags = re.sub(r"<[^>]+>", " ", without_hidden)
            compact = re.sub(r"\s+", " ", without_tags).strip()
            return compact or decoded.strip()
        soup = BeautifulSoup(decoded, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        return soup.get_text("\n", strip=True)

    def _image_to_data_url(self, mime_type: str, payload: bytes) -> str:
        encoded = base64.b64encode(payload).decode("ascii")
        return f"data:{mime_type};base64,{encoded}"
